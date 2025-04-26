from docx import Document
from weasyprint import HTML
import os

print("Convertir archivos word a pdf, SIN WORD")

def docx_to_html (docx_file):
    doc = Document(docx_file)
    html_content = "<html><body>"
    # Iterar sobre los párrafos del documento
    for para in doc.paragraphs:
        # Convertir cada párrafo a HTML
        html_content += f"<p>{para.text}</p>"
    
    # Iterar sobre las tablas del documento (opcional)
    for table in doc.tables:
        html_content += "<table border='1'>"
        for row in table.rows:
            html_content += "<tr>"
            for cell in row.cells:
                html_content += f"<td>{cell.text}</td>"
            html_content += "</tr>"
        html_content += "</table>"
    
    html_content += "</body></html>"
    return html_content

def html_to_pdf(html_content, pdf_file):
    # Generar el PDF desde el HTML
    HTML(string=html_content).write_pdf(pdf_file)

def convert_docx_to_pdf(docx_file, pdf_file):
    # Convertir el archivo .docx a HTML
    html_content = docx_to_html(docx_file)
    # Convertir el HTML a PDF
    html_to_pdf(html_content, pdf_file)
    print(f"Archivo convertido exitosamente: {pdf_file}")


docx_path = r"D:\Usuarios\Javier\Proyectos-Dev\Python - Aprendiendo\word-to-pdf\ARCHIVOWORD2.docx"
pdf_path = r"D:\Usuarios\Javier\Proyectos-Dev\Python - Aprendiendo\word-to-pdf\ARCHIVOWORD2.pdf"

convert_docx_to_pdf(docx_path, pdf_path)